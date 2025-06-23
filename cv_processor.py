import os
import re
import PyPDF2
import docx
from pathlib import Path
from typing import Dict, List, Optional, Union
import pandas as pd
from datetime import datetime
import json


class CVProcessor:
    """Advanced CV processing with support for multiple formats"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt']
        
    def process_cv(self, cv_path: str) -> Dict:
        """Process CV file and extract structured information"""
        file_path = Path(cv_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"CV file not found: {cv_path}")
        
        # Extract text based on file format
        if file_path.suffix.lower() == '.pdf':
            text = self._extract_pdf_text(cv_path)
        elif file_path.suffix.lower() == '.docx':
            text = self._extract_docx_text(cv_path)
        elif file_path.suffix.lower() == '.doc':
            text = self._extract_doc_text(cv_path)
        elif file_path.suffix.lower() == '.txt':
            with open(cv_path, 'r', encoding='utf-8') as f:
                text = f.read()
        else:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")
        
        # Extract structured information
        cv_data = {
            'file_path': cv_path,
            'raw_text': text,
            'contact_info': self._extract_contact_info(text),
            'experience': self._extract_experience(text),
            'education': self._extract_education(text),
            'skills': self._extract_skills(text),
            'certifications': self._extract_certifications(text),
            'languages': self._extract_languages(text),
            'summary': self._extract_summary(text),
            'total_experience_years': self._calculate_total_experience(text)
        }
        
        return cv_data
    
    def _extract_pdf_text(self, pdf_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            with open(pdf_path, 'rb') as file:
                # Try to read PDF with strict=False to handle corrupted PDFs
                try:
                    pdf_reader = PyPDF2.PdfReader(file, strict=False)
                except:
                    # If that fails, try with PyPDF2.PdfFileReader (older method)
                    file.seek(0)
                    pdf_reader = PyPDF2.PdfFileReader(file)
                
                # Get number of pages
                try:
                    num_pages = len(pdf_reader.pages)
                except:
                    num_pages = pdf_reader.getNumPages()
                
                # Extract text from each page
                for page_num in range(num_pages):
                    try:
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    except Exception as page_error:
                        continue
                        
        except Exception as e:
            pass
            
            # Try alternative method using PyPDF2 with different approach
            try:
                import io
                with open(pdf_path, 'rb') as file:
                    pdf_content = file.read()
                    pdf_file = io.BytesIO(pdf_content)
                    pdf_reader = PyPDF2.PdfReader(pdf_file, strict=False)
                    
                    for page in pdf_reader.pages:
                        try:
                            text += page.extract_text() + "\n"
                        except:
                            continue
                            
            except Exception as alt_error:
                pass
                
                # Try pdfplumber as last resort
                try:
                    import pdfplumber
                    with pdfplumber.open(pdf_path) as pdf:
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n"
                except Exception as plumber_error:
                    # Return empty string if all methods fail
                    text = ""
                
        # Clean up text
        text = text.strip()
        return text
    
    def _extract_docx_text(self, docx_path: str) -> str:
        """Extract text from DOCX file including tables and headers"""
        text = ""
        
        try:
            # First attempt with python-docx
            doc = docx.Document(docx_path)
            text_parts = []
            
            # Extract paragraph text
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract table text
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            # Extract header/footer text (if accessible)
            try:
                for section in doc.sections:
                    # Headers
                    if hasattr(section, 'header') and section.header:
                        for paragraph in section.header.paragraphs:
                            if paragraph.text.strip():
                                text_parts.append(paragraph.text)
                    
                    # Footers
                    if hasattr(section, 'footer') and section.footer:
                        for paragraph in section.footer.paragraphs:
                            if paragraph.text.strip():
                                text_parts.append(paragraph.text)
            except:
                # Skip headers/footers if there are issues accessing them
                pass
            
            text = "\n".join(text_parts)
            
        except Exception as e:
            pass
            
            # Try alternative method using zipfile for corrupted docx files
            try:
                import zipfile
                import xml.etree.ElementTree as ET
                
                pass
                
                with zipfile.ZipFile(docx_path, 'r') as zip_file:
                    # List available files in the archive
                    file_list = zip_file.namelist()
                    
                    # Try different possible document locations
                    possible_docs = ['word/document.xml', 'document.xml', 'content.xml']
                    doc_found = False
                    
                    for doc_path in possible_docs:
                        if doc_path in file_list:
                            try:
                                xml_content = zip_file.read(doc_path)
                                tree = ET.fromstring(xml_content)
                                
                                # Extract text from XML with multiple namespace attempts
                                namespaces_to_try = [
                                    {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'},
                                    {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                                     'office': 'urn:oasis:names:tc:opendocument:xmlns:office:1.0',
                                     'text': 'urn:oasis:names:tc:opendocument:xmlns:text:1.0'}
                                ]
                                
                                for namespaces in namespaces_to_try:
                                    try:
                                        text_elements = tree.findall('.//w:t', namespaces) or tree.findall('.//text:p', namespaces)
                                        if text_elements:
                                            alt_text_parts = [elem.text for elem in text_elements if elem.text]
                                            text = " ".join(alt_text_parts)
                                            if text:
                                                doc_found = True
                                                break
                                    except:
                                        continue
                                
                                if doc_found:
                                    break
                                    
                            except Exception as xml_error:
                                continue
                    
                    if not doc_found:
                        # Last resort: try to extract any readable text from any XML files
                        for file_name in file_list:
                            if file_name.endswith('.xml'):
                                try:
                                    xml_content = zip_file.read(file_name)
                                    # Simple text extraction without XML parsing
                                    content_str = xml_content.decode('utf-8', errors='ignore')
                                    # Extract text between XML tags using regex
                                    import re
                                    text_matches = re.findall(r'>([^<]+)<', content_str)
                                    extracted_text = ' '.join([match.strip() for match in text_matches if match.strip() and len(match.strip()) > 2])
                                    if len(extracted_text) > 50:  # Only use if we got meaningful text
                                        text = extracted_text
                                        break
                                except:
                                    continue
                        
            except Exception as alt_error:
                text = ""
        
        # Clean up text
        text = text.strip()
        return text
    
    def _extract_doc_text(self, doc_path: str) -> str:
        """Extract text from DOC file (older format)"""
        text = ""
        
        try:
            # First try with python-docx which sometimes works for .doc files
            text = self._extract_docx_text(doc_path)
            if text.strip():
                return text
        except:
            pass
        
        try:
            # Try using docx2txt which has better .doc support
            import docx2txt
            text = docx2txt.process(doc_path)
            if text.strip():
                return text
        except ImportError:
            pass
        except Exception as e:
            pass
        
        try:
            # Try using olefile for older .doc format
            import olefile
            import struct
            
            ole = olefile.OleFileIO(doc_path)
            
            # Try to find Word Document stream
            if ole.exists('WordDocument'):
                stream = ole.openstream('WordDocument')
                content = stream.read()
                
                # Extract text from the stream (basic extraction)
                # This is a simplified approach - real .doc parsing is complex
                extracted_parts = []
                
                # Try to find readable text in the binary content
                current_text = b''
                for byte in content:
                    if 32 <= byte <= 126:  # Printable ASCII range
                        current_text += bytes([byte])
                    else:
                        if len(current_text) > 3:  # Only keep text chunks longer than 3 chars
                            try:
                                decoded = current_text.decode('ascii', errors='ignore')
                                if decoded.strip():
                                    extracted_parts.append(decoded)
                            except:
                                pass
                        current_text = b''
                
                text = ' '.join(extracted_parts)
                ole.close()
                
                if text.strip():
                    return text
        except ImportError:
            pass
        except Exception as e:
            pass
        
        try:
            # Try using win32com if on Windows
            import platform
            if platform.system() == 'Windows':
                import win32com.client
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(doc_path)
                text = doc.Range().Text
                doc.Close()
                word.Quit()
                if text.strip():
                    return text
        except:
            pass
        
        try:
            # Try using subprocess with antiword (if installed)
            import subprocess
            result = subprocess.run(['antiword', doc_path], capture_output=True, text=True)
            if result.returncode == 0 and result.stdout.strip():
                text = result.stdout
                return text
        except:
            pass
        
        try:
            # Try using subprocess with catdoc (if installed)
            import subprocess
            result = subprocess.run(['catdoc', doc_path], capture_output=True, text=True)
            if result.returncode == 0 and result.stdout.strip():
                text = result.stdout
                return text
        except:
            pass
        
        try:
            # Try using subprocess with wvText (if installed)
            import subprocess
            result = subprocess.run(['wvText', doc_path, '-'], capture_output=True, text=True)
            if result.returncode == 0 and result.stdout.strip():
                text = result.stdout
                return text
        except:
            pass
        
        try:
            # Last resort: try to extract with textract
            import textract
            text = textract.process(doc_path).decode('utf-8')
            if text.strip():
                return text
        except:
            pass
        
        # If all methods fail, return empty string
        return ""
    
    def extract_name(self, text: str) -> str:
        """
        Extract candidate name from CV text using multiple strategies
        Handles various CV formats where name can appear anywhere
        """
        if not text or not text.strip():
            return ""
        
        # Clean text for processing
        text = text.strip()
        lines = text.split('\n')
        
        # Strategy 1: Look for explicit name patterns
        name_patterns = [
            # Name: John Doe or Name - John Doe
            r'(?:Name|NAME|Full Name|FULL NAME|Candidate Name)\s*[:|-]\s*([A-Za-z]+(?:\s+[A-Za-z]+)+)',
            # Personal Information section
            r'(?:Personal Information|PERSONAL INFORMATION|Personal Details)\s*\n+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)',
            # Contact Information followed by name
            r'(?:Contact Information|CONTACT INFORMATION|Contact Details)\s*\n+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)',
        ]
        
        for pattern in name_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                name = match.group(1).strip()
                # Validate name
                if self._is_valid_name(name):
                    return name
        
        # Strategy 2: Look at the first few non-empty lines
        # Most CVs start with the candidate's name
        for i, line in enumerate(lines[:10]):
            line = line.strip()
            if not line:
                continue
            
            # Skip common headers and non-name content
            skip_patterns = [
                r'curriculum\s*vitae', r'resume', r'cv\b', r'profile',
                r'@', r'http', r'www\.', r'\.com', r'page\s*\d+',
                r'confidential', r'references', r'objective', r'summary'
            ]
            
            if any(re.search(pattern, line, re.IGNORECASE) for pattern in skip_patterns):
                continue
            
            # Skip lines with too many numbers (likely phone/address)
            if len(re.findall(r'\d', line)) > 5:
                continue
            
            # Check if line looks like a name
            if self._looks_like_name(line):
                return line
        
        # Strategy 3: Find name near email address
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_match = re.search(email_pattern, text)
        
        if email_match:
            email = email_match.group()
            email_pos = text.find(email)
            
            # Extract context around email (300 chars before and after)
            context_start = max(0, email_pos - 300)
            context_end = min(len(text), email_pos + 300)
            context = text[context_start:context_end]
            context_lines = context.split('\n')
            
            # Look for name-like patterns near email
            for line in context_lines:
                line = line.strip()
                if line and email not in line and self._looks_like_name(line):
                    return line
            
            # Try to extract name from email prefix
            email_prefix = email.split('@')[0]
            # Remove numbers and common separators
            name_from_email = re.sub(r'[\d._-]', ' ', email_prefix).strip()
            
            # Search for this name pattern in the document
            if name_from_email:
                words = name_from_email.split()
                if words:
                    # Look for these words appearing together in proper case
                    for i in range(len(lines)):
                        line = lines[i].strip()
                        if all(word.lower() in line.lower() for word in words):
                            # Check if this line is a properly formatted name
                            if self._looks_like_name(line):
                                return line
        
        # Strategy 4: Look for the largest text or bold text in first page
        # This requires looking for patterns that might indicate emphasis
        emphasized_patterns = [
            r'^([A-Z\s]+)$',  # All caps name
            r'^([A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+)+)\s*$',  # Title case full name
        ]
        
        for i, line in enumerate(lines[:20]):
            line = line.strip()
            if not line or len(line) > 50:  # Skip long lines
                continue
                
            for pattern in emphasized_patterns:
                match = re.match(pattern, line)
                if match:
                    potential_name = match.group(1).strip()
                    if self._is_valid_name(potential_name):
                        return potential_name.title()  # Convert to proper case
        
        # Strategy 5: Machine learning based approach using common name indicators
        # Look for lines that have name-like characteristics
        name_indicators = {
            'length': lambda x: 10 <= len(x) <= 40,
            'words': lambda x: 2 <= len(x.split()) <= 4,
            'capitals': lambda x: sum(1 for c in x if c.isupper()) >= 2,
            'no_special': lambda x: not re.search(r'[!@#$%^&*()_+=\[\]{};:"\\|,.<>?/]', x),
            'has_letters': lambda x: sum(1 for c in x if c.isalpha()) >= 5,
        }
        
        candidates = []
        for i, line in enumerate(lines[:30]):
            line = line.strip()
            if not line:
                continue
                
            score = sum(1 for check in name_indicators.values() if check(line))
            if score >= 4:
                candidates.append((score, i, line))
        
        # Sort by score and position (prefer higher score and earlier position)
        candidates.sort(key=lambda x: (-x[0], x[1]))
        
        for score, pos, candidate in candidates[:5]:
            if self._is_valid_name(candidate):
                return candidate
        
        return ""  # Return empty string if no name found
    
    def _looks_like_name(self, text: str) -> bool:
        """Check if text looks like a person's name"""
        if not text:
            return False
            
        # Remove extra spaces
        text = ' '.join(text.split())
        
        # Check length
        if len(text) < 3 or len(text) > 50:
            return False
        
        # Check word count (names typically have 2-4 words)
        words = text.split()
        if len(words) < 1 or len(words) > 5:
            return False
        
        # Check if mostly alphabetic
        alpha_chars = sum(1 for c in text if c.isalpha() or c.isspace())
        if alpha_chars < len(text) * 0.8:
            return False
        
        # Check for common name patterns
        # At least first letter of each word should be capital or all could be capital
        is_title_case = all(word[0].isupper() or len(word) <= 2 for word in words if word)
        is_all_caps = text.isupper()
        
        if not (is_title_case or is_all_caps):
            return False
        
        # Exclude common false positives
        false_positives = [
            'curriculum vitae', 'resume', 'personal information', 'contact details',
            'professional summary', 'work experience', 'education', 'skills',
            'references', 'page', 'confidential', 'private', 'address', 'email',
            'phone', 'mobile', 'linkedin', 'objective', 'profile', 'summary'
        ]
        
        text_lower = text.lower()
        if any(fp in text_lower for fp in false_positives):
            return False
        
        return True
    
    def _is_valid_name(self, name: str) -> bool:
        """Validate if extracted text is a valid name"""
        if not name:
            return False
            
        # Clean the name
        name = ' '.join(name.split())
        
        # Check basic criteria
        if len(name) < 3 or len(name) > 50:
            return False
            
        words = name.split()
        if len(words) < 2:  # Require at least first and last name
            return False
            
        # Check if each word is valid
        for word in words:
            # Allow short words like "Jr", "Sr", "de", "van"
            if len(word) <= 2:
                continue
            # Check if word contains mostly letters
            if not word.replace('-', '').replace("'", '').isalpha():
                return False
        
        return True
    
    def _extract_contact_info(self, text: str) -> Dict:
        """Extract contact information"""
        contact = {}
        
        # Email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_match = re.search(email_pattern, text)
        if email_match:
            contact['email'] = email_match.group()
        
        # Phone
        phone_patterns = [
            r'(?:\+?61|0)[2-478](?:[ -]?[0-9]){8}',  # Australian format
            r'(?:\+?1-)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',  # US format
            r'\+?[0-9]{10,15}'  # International
        ]
        for pattern in phone_patterns:
            phone_match = re.search(pattern, text)
            if phone_match:
                contact['phone'] = phone_match.group()
                break
        
        # LinkedIn
        linkedin_pattern = r'(?:linkedin\.com/in/|linkedin:)\s*([^\s]+)'
        linkedin_match = re.search(linkedin_pattern, text, re.IGNORECASE)
        if linkedin_match:
            contact['linkedin'] = linkedin_match.group(1)
        
        # Location
        location_patterns = [
            r'(?:Location|Address|Based in):\s*([^\n]+)',
            r'([A-Za-z\s]+,\s*[A-Z]{2,3}(?:\s*\d{4,5})?)'  # City, State/Country
        ]
        for pattern in location_patterns:
            location_match = re.search(pattern, text, re.IGNORECASE)
            if location_match:
                contact['location'] = location_match.group(1).strip()
                break
        
        return contact
    
    def _extract_experience(self, text: str) -> List[Dict]:
        """Extract work experience"""
        experiences = []
        
        # Common section headers
        exp_headers = [
            r'(?:work\s*)?experience',
            r'employment\s*history',
            r'professional\s*experience',
            r'career\s*history'
        ]
        
        # Find experience section
        exp_section = None
        for header in exp_headers:
            pattern = rf'(?i)(?:^|\n)({header})[:\s]*\n(.*?)(?=\n(?:education|skills|certifications|references|$))'
            match = re.search(pattern, text, re.DOTALL)
            if match:
                exp_section = match.group(2)
                break
        
        if exp_section:
            # Extract individual experiences
            # Pattern for job entries (Company, Position, Dates)
            job_pattern = r'([A-Za-z\s\.,&]+?)\s*[-–|]\s*([A-Za-z\s]+?)\s*\n?\s*\(?((?:\d{4}|\d{2}/\d{2}/\d{4}|[A-Za-z]+\s*\d{4}))(?:\s*[-–]\s*)?(?:(?:\d{4}|\d{2}/\d{2}/\d{4}|[A-Za-z]+\s*\d{4})|Present|Current)?\)?'
            
            for match in re.finditer(job_pattern, exp_section):
                exp = {
                    'company': match.group(1).strip(),
                    'position': match.group(2).strip(),
                    'dates': match.group(3).strip() if match.group(3) else ""
                }
                experiences.append(exp)
        
        return experiences
    
    def _extract_education(self, text: str) -> List[Dict]:
        """Extract education information"""
        education = []
        
        # Education section headers
        edu_headers = [
            r'education',
            r'academic\s*background',
            r'qualifications'
        ]
        
        # Find education section
        edu_section = None
        for header in edu_headers:
            pattern = rf'(?i)(?:^|\n)({header})[:\s]*\n(.*?)(?=\n(?:experience|skills|certifications|references|$))'
            match = re.search(pattern, text, re.DOTALL)
            if match:
                edu_section = match.group(2)
                break
        
        # Common degree patterns
        degree_patterns = [
            r'(Bachelor|Master|PhD|Ph\.D\.|Doctorate|MBA|Diploma|Certificate)(?:\'s|\s*of)?\s*(?:in\s*)?([A-Za-z\s,]+?)(?:\s*[-–]\s*)?([A-Za-z\s]+(?:University|College|Institute|School))',
            r'([A-Za-z\s]+(?:University|College|Institute))\s*[-–]\s*(Bachelor|Master|PhD|MBA|Diploma)',
        ]
        
        if edu_section:
            for pattern in degree_patterns:
                for match in re.finditer(pattern, edu_section, re.IGNORECASE):
                    edu = {
                        'degree': match.group(1),
                        'field': match.group(2) if len(match.groups()) > 2 else "",
                        'institution': match.group(3) if len(match.groups()) > 2 else match.group(1)
                    }
                    education.append(edu)
        
        return education
    
    def _extract_skills(self, text: str) -> List[str]:
        """Extract skills"""
        skills = []
        
        # Skills section headers
        skill_headers = [
            r'skills',
            r'core\s*competencies',
            r'technical\s*skills',
            r'key\s*skills'
        ]
        
        # Find skills section
        skills_section = None
        for header in skill_headers:
            pattern = rf'(?i)(?:^|\n)({header})[:\s]*\n(.*?)(?=\n(?:experience|education|certifications|references|$))'
            match = re.search(pattern, text, re.DOTALL)
            if match:
                skills_section = match.group(2)
                break
        
        if skills_section:
            # Extract skills (comma or bullet separated)
            skill_lines = skills_section.split('\n')
            for line in skill_lines:
                # Remove bullets and clean
                clean_line = re.sub(r'^[\s•·-]+', '', line).strip()
                if clean_line and len(clean_line) < 100:  # Avoid long descriptions
                    # Split by commas if present
                    if ',' in clean_line:
                        skills.extend([s.strip() for s in clean_line.split(',') if s.strip()])
                    else:
                        skills.append(clean_line)
        
        # Also look for inline skills throughout the document
        skill_keywords = [
            'leadership', 'management', 'strategic planning', 'project management',
            'stakeholder management', 'business development', 'financial management',
            'negotiation', 'communication', 'team building', 'problem solving',
            'analytical', 'customer service', 'sales', 'marketing', 'operations',
            'change management', 'risk management', 'compliance', 'governance'
        ]
        
        text_lower = text.lower()
        for skill in skill_keywords:
            if skill in text_lower and skill not in [s.lower() for s in skills]:
                skills.append(skill.title())
        
        return list(set(skills))[:20]  # Return top 20 unique skills
    
    def _extract_certifications(self, text: str) -> List[str]:
        """Extract certifications"""
        certifications = []
        
        # Certification patterns
        cert_patterns = [
            r'(?:certified|certification|certificate)\s+(?:in\s+)?([A-Za-z\s\(\)]+)',
            r'([A-Z]{2,}\s*[A-Z]*)\s*(?:certified|certification)',
            r'(?:PMP|PRINCE2|Six Sigma|ITIL|Agile|Scrum|AWS|Azure|Google Cloud)'
        ]
        
        for pattern in cert_patterns:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                cert = match.group(1) if match.groups() else match.group()
                if cert and len(cert) < 50:
                    certifications.append(cert.strip())
        
        return list(set(certifications))
    
    def _extract_languages(self, text: str) -> List[str]:
        """Extract languages"""
        languages = []
        
        # Language section pattern
        lang_pattern = r'(?i)languages?[:\s]+([^\n]+(?:\n[^\n]+)*?)(?=\n(?:[A-Z]|$))'
        match = re.search(lang_pattern, text)
        
        if match:
            lang_text = match.group(1)
            # Common language names
            common_languages = [
                'English', 'Spanish', 'French', 'German', 'Chinese', 'Mandarin',
                'Cantonese', 'Japanese', 'Korean', 'Arabic', 'Hindi', 'Portuguese',
                'Russian', 'Italian', 'Dutch'
            ]
            
            for lang in common_languages:
                if lang.lower() in lang_text.lower():
                    languages.append(lang)
        
        return languages
    
    def _extract_summary(self, text: str) -> str:
        """Extract professional summary"""
        summary = ""
        
        # Summary section headers
        summary_headers = [
            r'(?:professional\s*)?summary',
            r'(?:executive\s*)?profile',
            r'about\s*me',
            r'objective'
        ]
        
        for header in summary_headers:
            pattern = rf'(?i)(?:^|\n)({header})[:\s]*\n(.*?)(?=\n(?:[A-Z]{{2,}}|experience|education|skills|$))'
            match = re.search(pattern, text, re.DOTALL)
            if match:
                summary = match.group(2).strip()
                # Clean up and limit length
                summary = ' '.join(summary.split())[:500]
                break
        
        return summary
    
    def _calculate_total_experience(self, text: str) -> int:
        """Calculate total years of experience"""
        # Look for explicit statements
        exp_patterns = [
            r'(\d+)\+?\s*years?\s*(?:of\s*)?(?:total\s*)?experience',
            r'(?:over|more than)\s*(\d+)\s*years?',
            r'(\d+)\s*years?\s*(?:in|as|of)'
        ]
        
        max_years = 0
        for pattern in exp_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                years = int(match.group(1))
                max_years = max(max_years, years)
        
        # If no explicit statement, try to calculate from work history
        if max_years == 0:
            # Look for date ranges
            date_pattern = r'(\d{4})\s*[-–]\s*(?:(\d{4})|present|current)'
            matches = re.finditer(date_pattern, text, re.IGNORECASE)
            
            earliest_year = datetime.now().year
            latest_year = 0
            
            for match in matches:
                start_year = int(match.group(1))
                end_year = int(match.group(2)) if match.group(2) and match.group(2).isdigit() else datetime.now().year
                
                earliest_year = min(earliest_year, start_year)
                latest_year = max(latest_year, end_year)
            
            if latest_year > 0:
                max_years = latest_year - earliest_year
        
        return max_years


class BatchCVProcessor:
    """Process multiple CVs and prepare data for training"""
    
    def __init__(self):
        self.cv_processor = CVProcessor()
        
    def process_job_folder(self, job_folder_path: str) -> pd.DataFrame:
        """Process all CVs in a job folder"""
        job_folder = Path(job_folder_path)
        
        if not job_folder.exists():
            raise ValueError(f"Job folder not found: {job_folder_path}")
        
        # Read job description
        job_desc_path = job_folder / "job_description.txt"
        if not job_desc_path.exists():
            raise ValueError(f"Job description not found in {job_folder_path}")
        
        with open(job_desc_path, 'r', encoding='utf-8') as f:
            job_description = f.read()
        
        # Process CVs from each stage
        all_cv_data = []
        files_processed = 0
        files_skipped = 0
        files_failed = 0
        
        for stage in ['REJECT', 'SHORTLIST', 'INTERVIEW', 'ACCEPT']:
            stage_folder = job_folder / stage
            if not stage_folder.exists():
                continue
            
            for cv_file in stage_folder.iterdir():
                if cv_file.suffix.lower() in self.cv_processor.supported_formats:
                    # Basic file validation
                    if not cv_file.is_file() or cv_file.stat().st_size == 0:
                        files_skipped += 1
                        continue
                    
                    try:
                        cv_data = self.cv_processor.process_cv(str(cv_file))
                        
                        # Only add if we extracted some meaningful text
                        if cv_data.get('raw_text', '').strip():
                            cv_data['stage'] = stage
                            cv_data['job_id'] = job_folder.name
                            cv_data['job_description'] = job_description
                            all_cv_data.append(cv_data)
                            files_processed += 1
                        else:
                            files_skipped += 1
                            
                    except Exception as e:
                        files_failed += 1
        
        # Processing summary
        total_files = files_processed + files_skipped + files_failed
        
        # Convert to DataFrame
        df = pd.DataFrame(all_cv_data)
        return df
    
    def process_all_jobs(self, jobs_directory: str) -> pd.DataFrame:
        """Process all job folders"""
        jobs_path = Path(jobs_directory)
        all_data = []
        
        for job_folder in jobs_path.iterdir():
            if job_folder.is_dir():
                try:
                    df = self.process_job_folder(str(job_folder))
                    all_data.append(df)
                except Exception as e:
                    pass
        
        if all_data:
            return pd.concat(all_data, ignore_index=True)
        else:
            return pd.DataFrame()


if __name__ == "__main__":
    # Example usage
    processor = CVProcessor()
    
    # Process a single CV
    # cv_data = processor.process_cv("path/to/cv.pdf")
    # print(json.dumps(cv_data, indent=2))
    
    # Batch process
    batch_processor = BatchCVProcessor()
    # df = batch_processor.process_all_jobs("jobs")
    # print(f"Processed {len(df)} CVs")
    # df.to_csv("processed_cvs.csv", index=False)